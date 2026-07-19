"""Regression test for the docx export 500 on deployed dev (QA round 1, Finding 1).

Root cause: `Dockerfile.prod` installs `requirements.txt` (the lean list used
for the deployed container), while local/native dev installs
`requirements-macos.txt`. The latter has always listed `python-docx`, but the
former did not, so `synapsis/exporters/docx.py`'s `from docx import Document`
raised ImportError -> HTTP 500 in the deployed dev container even though the
endpoint worked locally.

These tests are cheap and catch both halves of the regression:
  1. `docx` must actually import in whatever environment tests are run in
     (a hard assert, not `pytest.importorskip`, since we DO require it now).
  2. `requirements.txt` — the file Dockerfile.prod installs — must pin
     `python-docx`, so this can never silently drift again.
"""

from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent


def test_docx_module_imports_cleanly():
    """The docx exporter's core dependency must be importable.

    synapsis/exporters/docx.py guards this import and raises HTTPException(500)
    if it's missing — this test asserts that guard never trips.
    """
    import docx  # noqa: F401
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401

    # Sanity: a real Document object can be constructed (exercises the C-ish
    # lxml plumbing beyond a bare import).
    doc = DocxDocument()
    doc.add_paragraph("smoke test")
    assert doc.paragraphs[0].text == "smoke test"


def test_requirements_txt_pins_python_docx():
    """`requirements.txt` (installed by Dockerfile.prod into the deployed
    dev/prod image) must list python-docx — this is what the previously
    deployed image was missing.
    """
    requirements_txt = (REPO_ROOT / "requirements.txt").read_text()
    lines = [
        line.strip()
        for line in requirements_txt.splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    docx_lines = [line for line in lines if line.lower().startswith("python-docx")]
    assert docx_lines, (
        "requirements.txt must pin python-docx (Dockerfile.prod installs this "
        "file into the deployed container; synapsis/exporters/docx.py requires it)."
    )


def test_requirements_macos_also_pins_python_docx():
    """Guard the other direction too: keep both requirement files in sync for
    this dependency so local dev and the deployed container never diverge again.
    """
    requirements_macos = (REPO_ROOT / "requirements-macos.txt").read_text()
    assert any(
        line.strip().lower().startswith("python-docx")
        for line in requirements_macos.splitlines()
    )
