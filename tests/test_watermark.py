"""Tests for the shared AI-content watermark utility and its application to
every export format (Step 1 of the July-7 guardrails sprint).

The mandate (Jules's standing AI-disclaimer requirement + CGIAR SO SOP) is that
EVERY export the tool produces carries the standardized AI-content notice. These
tests assert the notice is present in each supported format's output.
"""

import io
import re
from datetime import datetime
from pathlib import Path

import pytest

from synapsis.exporters import watermark as wm

REPO_ROOT = Path(__file__).resolve().parent.parent
CONTACTS_TS = REPO_ROOT / "frontend/src/components/guardrails/contacts.ts"
WATERMARK_PY = Path(wm.__file__)

#: A pinned export moment so timestamp-bearing assertions stay deterministic.
PINNED = datetime(2026, 7, 20, 14, 30)


# ---------------------------------------------------------------------------
# Core wording
# ---------------------------------------------------------------------------

def test_banner_and_notice_wording():
    assert "AI V0 DRAFT" in wm.WATERMARK_BANNER
    notice = wm.provenance_notice(datetime(2026, 7, 20))
    assert "CGIAR innovation data (PRMS)" in notice
    assert "AI-added interpretation" in notice
    assert "human quality assurance" in notice
    assert "Data as of 2026-07-20." in notice


def test_sop_disclosure_present():
    # The verbatim CGIAR SO SOP disclosure sentence must be reused.
    assert "reviewed, validated, and approved by responsible human authors" in wm.SOP_DISCLOSURE


def test_markdown_watermark_block():
    md = wm.watermark_markdown(PINNED)
    assert md.startswith("> ")
    assert wm.WATERMARK_BANNER in md
    assert "human quality assurance" in md
    # Contact route + export timestamp ride in the same blockquote.
    assert wm.contact_line_text() in md
    assert wm.export_timestamp_line(PINNED) in md
    # Every line of the block stays inside the blockquote.
    assert all(line.startswith(">") for line in md.strip().splitlines())


def test_html_watermark_block_and_css():
    html = wm.watermark_html(PINNED)
    assert 'class="ai-watermark"' in html
    assert wm.WATERMARK_BANNER in html
    assert "ai-watermark" in wm.WATERMARK_HTML_CSS
    assert 'class="ai-watermark-contact"' in html
    assert 'class="ai-watermark-timestamp"' in html
    assert wm.export_timestamp_line(PINNED) in html


def test_docx_watermark_applied_at_top_and_footer():
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    doc.add_paragraph("Some body content.")
    wm.apply_ai_watermark(doc, date=PINNED)

    texts = [p.text for p in doc.paragraphs]
    # Banner is the very first paragraph.
    assert texts[0] == wm.WATERMARK_BANNER
    # Provenance + SOP + contact line follow, ahead of the body.
    assert any("CGIAR innovation data (PRMS)" in t for t in texts[:4])
    assert any("responsible human authors" in t for t in texts[:4])
    assert any("Reach out before you use it" in t for t in texts[:4])
    assert texts[4] == "Some body content."
    # Footer carries the banner on every page.
    footer_text = " ".join(p.text for p in doc.sections[0].footer.paragraphs)
    assert wm.WATERMARK_BANNER in footer_text


# ---------------------------------------------------------------------------
# Contact route (guardrails "reach out if in doubt")
# ---------------------------------------------------------------------------

def test_contact_line_text_wording():
    line = wm.contact_line_text()
    assert line.startswith(wm.CONTACT_LEAD_IN)
    assert "Marc Schut (scope & use, marc.schut@cgiar.org)" in line
    assert "Jose Luis Berenguer (technical, jose@synapsis-analytics.com)" in line
    assert " or " in line
    assert line.endswith(".")


def test_contact_line_html_has_mailto_links_and_escapes():
    html = wm.contact_line_html()
    for _name, email, _remit in wm.GUARDRAIL_CONTACTS:
        assert f'<a href="mailto:{email}">{email}</a>' in html
    # The ampersand in the "scope & use" remit must be HTML-escaped.
    assert "scope &amp; use" in html
    assert "scope & use" not in html


def test_contacts_do_not_drift_from_frontend_master():
    """`contacts.ts` is the master copy — the Python mirror must match verbatim.

    Parses the TypeScript source and asserts every name, email, remit and the
    lead-in sentence appear both in the Python constants and literally in
    watermark.py, so the two copies can never silently diverge.
    """
    assert CONTACTS_TS.is_file(), f"master contact file missing: {CONTACTS_TS}"
    ts_src = CONTACTS_TS.read_text(encoding="utf-8")
    py_src = WATERMARK_PY.read_text(encoding="utf-8")

    # --- contacts array ---
    # NB: split on the array literal's own "= [", not on the first "]" after
    # the identifier — the type annotation `GuardrailContact[]` has one too.
    array_block = re.search(
        r"GUARDRAIL_CONTACTS[^=]*=\s*\[(?P<items>.*?)\n\]", ts_src, re.DOTALL
    ).group("items")
    ts_contacts = [
        (m.group("name"), m.group("email"), m.group("remit"))
        for m in re.finditer(
            r"name:\s*'(?P<name>[^']+)'\s*,\s*"
            r"email:\s*'(?P<email>[^']+)'\s*,\s*"
            r"remit:\s*'(?P<remit>[^']+)'",
            re.sub(r"\s+", " ", array_block),
        )
    ]
    assert ts_contacts, "could not parse GUARDRAIL_CONTACTS out of contacts.ts"
    assert ts_contacts == wm.GUARDRAIL_CONTACTS, (
        "synapsis/exporters/watermark.py has drifted from its master copy "
        "frontend/src/components/guardrails/contacts.ts — update both."
    )
    for name, email, remit in ts_contacts:
        assert name in py_src and email in py_src and remit in py_src

    # --- lead-in sentence ---
    lead_in = re.search(r"CONTACT_LEAD_IN\s*=\s*'([^']+)'", ts_src).group(1)
    assert lead_in == wm.CONTACT_LEAD_IN
    assert lead_in in py_src


def test_contact_line_matches_frontend_sentence_plus_emails():
    """The export sentence is the in-app CONTACT_LINE_TEXT with emails added.

    Exports travel outside the app, where mailto links and the UI contact route
    are unavailable, so the addresses have to be readable on the page. The rest
    of the sentence must be identical to the frontend's.
    """
    ts_src = CONTACTS_TS.read_text(encoding="utf-8")
    lead_in = re.search(r"CONTACT_LEAD_IN\s*=\s*'([^']+)'", ts_src).group(1)
    # Frontend renders "Name (remit)"; the export adds ", email" inside the parens.
    frontend_sentence = f"{lead_in} " + " or ".join(
        f"{n} ({r})" for n, _e, r in wm.GUARDRAIL_CONTACTS
    ) + "."
    exported = wm.contact_line_text()
    stripped = re.sub(r", [^()]+@[^()]+\)", ")", exported)
    assert stripped == frontend_sentence


# ---------------------------------------------------------------------------
# Export timestamp
# ---------------------------------------------------------------------------

def test_export_timestamp_line_is_pinnable_and_utc():
    line = wm.export_timestamp_line(datetime(2026, 8, 9, 14, 5))
    assert line == "Export generated on 9 August 2026 at 14:05 UTC."
    # No zero-padded day (matches the donor's en-GB "9 August 2026" rendering).
    assert "09 August" not in line


def test_export_timestamp_line_defaults_to_now_utc():
    line = wm.export_timestamp_line()
    assert line.startswith("Export generated on ")
    assert line.endswith(" UTC.")


def test_provenance_notice_is_not_reworded_by_the_timestamp():
    # The signed-off provenance wording must stay untouched.
    assert wm.PROVENANCE_NOTICE.endswith("Data as of {date}.")
    assert "Export generated" not in wm.PROVENANCE_NOTICE


# ---------------------------------------------------------------------------
# HTML / PDF per-page overlay
# ---------------------------------------------------------------------------

def test_watermark_html_overlay_structure():
    overlay = wm.watermark_html_overlay(PINNED)
    assert 'class="ai-watermark-overlay"' in overlay
    assert 'class="ai-watermark-pagefooter"' in overlay
    assert overlay.count(wm.WATERMARK_BANNER) == 2
    assert wm.PRODUCT_FOOTER in overlay
    assert wm.provenance_notice(PINNED) in overlay
    # Decorative furniture must be hidden from assistive tech.
    assert overlay.count('aria-hidden="true"') == 2


def test_watermark_overlay_css_rules_present():
    css = wm.WATERMARK_HTML_CSS
    assert ".ai-watermark-overlay{" in css
    assert "position:fixed" in css
    assert "rotate(-45deg)" in css
    assert "pointer-events:none" in css
    assert "user-select:none" in css
    assert "opacity:.08" in css
    assert "@media print{" in css
    # The overlay is a full-page flex container, not a top/left:50% offset —
    # only the former repeats on every printed page in Chromium.
    assert ".ai-watermark-overlay{position:fixed;top:0;left:0;width:100%;height:100%;" in css
    assert "display:flex;align-items:center;justify-content:center;" in css


def test_overlay_and_footer_layering():
    """The marks paint above content, and the footer above the overlay.

    IA's exports are built from opaque message cards, so a behind-content
    watermark is painted over and survives only in the page margins (observed
    in a headless-Chrome print render). The donor does the same: its PDF draws
    the draft TEXT watermark as the top layer.
    """
    css = wm.WATERMARK_HTML_CSS
    overlay_rule = css.split(".ai-watermark-overlay{", 1)[1].split("}", 1)[0]
    footer_rule = css.split("@media print{", 1)[1].split(
        ".ai-watermark-pagefooter{", 1
    )[1].split("}", 1)[0]

    overlay_z = int(re.search(r"z-index:(\d+)", overlay_rule).group(1))
    footer_z = int(re.search(r"z-index:(\d+)", footer_rule).group(1))
    assert footer_z > overlay_z, "print footer must sit above the diagonal overlay"

    # An on-top overlay must never intercept clicks or selection.
    assert "pointer-events:none" in overlay_rule
    assert "user-select:none" in overlay_rule
    # No blanket body>* stacking rule: it out-specified the footer's own
    # `position:fixed` and dropped it into normal flow at the top of page 1.
    assert "body>*" not in css


def test_watermark_plain_carries_contact_route():
    plain = wm.watermark_plain(PINNED)
    assert wm.WATERMARK_BANNER in plain
    assert wm.contact_line_text() in plain
    assert "\n" not in plain  # still a single line


# ---------------------------------------------------------------------------
# Markdown end-of-document footer
# ---------------------------------------------------------------------------

def test_watermark_markdown_footer():
    footer = wm.watermark_markdown_footer(PINNED)
    assert footer.startswith("\n---\n")
    assert wm.PRODUCT_FOOTER in footer
    assert wm.WATERMARK_BANNER in footer
    assert wm.provenance_notice(PINNED) in footer


# ---------------------------------------------------------------------------
# DOCX per-page marks (VML watermark, running header, page numbers, box)
# ---------------------------------------------------------------------------

def _roundtripped_doc(**kwargs):
    """Build a watermarked document, save it and reopen it from bytes.

    The save/reopen round-trip is the point: raw VML injected into the header
    must survive python-docx serialization to be worth anything in Word.
    """
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    doc.add_paragraph("Some body content.")
    wm.apply_ai_watermark(doc, date=PINNED, **kwargs)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return docx.Document(buf)


def test_docx_vml_watermark_survives_save_and_reopen():
    doc = _roundtripped_doc(title="Innovation Portfolio")
    header_xml = doc.sections[0].header._element.xml

    # The classic Word text watermark shape, in the header (= every page).
    assert "PowerPlusWaterMarkObject1" in header_xml
    assert "_x0000_t136" in header_xml  # text-on-path shapetype
    assert 'urn:schemas-microsoft-com:vml' in header_xml
    # Diagonal, red, faint, behind the text layer, page-anchored.
    assert "rotation:315" in header_xml
    assert 'fillcolor="#b40000"' in header_xml
    assert 'opacity=".13"' in header_xml
    assert "z-index:-251654144" in header_xml
    assert "mso-position-horizontal:center" in header_xml
    # The banner itself is the watermark text.
    assert "<v:textpath" in header_xml
    assert "AI V0 DRAFT" in header_xml


def test_docx_running_header_has_title_and_banner():
    doc = _roundtripped_doc(title="Innovation Portfolio")
    header = doc.sections[0].header
    header_text = " ".join(p.text for p in header.paragraphs)
    assert "Innovation Portfolio" in header_text
    assert wm.WATERMARK_BANNER in header_text
    # Thin rule under the running header.
    assert "w:pBdr" in header._element.xml


def test_docx_running_header_falls_back_to_product_name():
    doc = _roundtripped_doc()
    header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
    assert wm.PRODUCT_FOOTER in header_text
    assert wm.WATERMARK_BANNER in header_text


def test_docx_footer_has_page_number_field():
    doc = _roundtripped_doc()
    footer = doc.sections[0].footer
    footer_xml = footer._element.xml
    # A real PAGE field, not a literal number.
    assert "PAGE" in footer_xml
    assert 'w:fldCharType="begin"' in footer_xml
    assert 'w:fldCharType="end"' in footer_xml
    assert "Page " in footer_xml
    # Right tab stop so the number sits on the outer edge.
    assert 'w:val="right"' in footer_xml
    # Existing footer wording is preserved verbatim.
    footer_text = " ".join(p.text for p in footer.paragraphs)
    assert wm.PRODUCT_FOOTER in footer_text
    assert wm.WATERMARK_BANNER in footer_text
    assert wm.provenance_notice(PINNED) in footer_text


def test_docx_notice_box_is_bordered_and_shaded():
    from docx.oxml.ns import qn
    from lxml import etree

    doc = _roundtripped_doc()
    notice_paras = doc.paragraphs[:4]
    assert notice_paras[0].text == wm.WATERMARK_BANNER

    borders = []
    for para in notice_paras:
        pPr = para._p.get_or_add_pPr()
        pPr_xml = pPr.xml
        assert "w:pBdr" in pPr_xml, f"no border on notice paragraph: {para.text[:40]}"
        assert 'w:fill="FFF5F5"' in pPr_xml, "notice paragraph is not shaded"
        assert 'w:color="B40000"' in pPr_xml
        pBdr = pPr.find(qn("w:pBdr"))
        # All four edges on every paragraph — see _notice_box_pr_xml.
        edges = {child.tag for child in pBdr}
        assert edges == {qn(f"w:{s}") for s in ("top", "left", "bottom", "right")}
        # NB: w:pBdr has no python-docx element class, so no `.xml` property.
        borders.append(etree.tostring(pBdr))

    # THE invariant: Word/LibreOffice only merge consecutive paragraph borders
    # into one box when the definitions are identical. If these ever diverge,
    # the box renders as disconnected per-paragraph segments.
    assert len(set(borders)) == 1, "notice paragraph borders differ — box will not merge"


def test_docx_notice_box_carries_timestamp_and_contact():
    doc = _roundtripped_doc()
    texts = [p.text for p in doc.paragraphs[:4]]
    # Timestamp shares the provenance paragraph (zero-draft `break: 1` pattern).
    assert wm.export_timestamp_line(PINNED) in texts[1]
    assert wm.provenance_notice(PINNED) in texts[1]
    assert texts[3] == wm.contact_line_text()


def test_docx_pPr_children_are_in_schema_order():
    """Word rejects a `w:pPr` whose children are out of ECMA-376 order."""
    from docx.oxml.ns import qn

    doc = _roundtripped_doc(title="Ordering check")
    order = {qn(tag): i for i, tag in enumerate(wm._PPR_TAG_SEQ)}

    paragraphs = list(doc.paragraphs)
    paragraphs += list(doc.sections[0].header.paragraphs)
    paragraphs += list(doc.sections[0].footer.paragraphs)

    for para in paragraphs:
        pPr = para._p.pPr
        if pPr is None:
            continue
        ranks = [order[child.tag] for child in pPr if child.tag in order]
        assert ranks == sorted(ranks), (
            f"pPr children out of schema order in {para.text[:40]!r}: {ranks}"
        )


def test_apply_ai_watermark_is_no_op_safe_on_a_broken_document():
    """Header/footer/VML failures are swallowed; the notice box always lands."""
    docx = pytest.importorskip("docx")

    def _explode(self):
        raise RuntimeError("section access is broken in this document")

    doc = docx.Document()
    doc.add_paragraph("Body.")
    document_cls = type(doc)
    original = document_cls.sections
    try:
        document_cls.sections = property(_explode)
        wm.apply_ai_watermark(doc, date=PINNED)
    finally:
        document_cls.sections = original

    # Top-of-document notice still applied despite the header/footer explosion.
    assert doc.paragraphs[0].text == wm.WATERMARK_BANNER
    assert doc.paragraphs[3].text == wm.contact_line_text()


# ---------------------------------------------------------------------------
# End-to-end: each exporter emits the watermark
# ---------------------------------------------------------------------------

def _fake_rows():
    """Minimal DB-row stand-ins understood by the exporters (dict-like access)."""
    import json

    return [
        {"type": "user", "data": json.dumps({"content": "How many innovations?"}), "ts": 1_700_000_000.0},
        {"type": "text", "data": json.dumps({"content": "There are 2,755."}), "ts": 1_700_000_001.0},
    ]


def test_markdown_export_is_watermarked():
    from synapsis.exporters import export_markdown

    content, _ = export_markdown("Test", "abc123", _fake_rows(), "standard")
    assert wm.WATERMARK_BANNER in content
    assert "human quality assurance" in content
    assert wm.contact_line_text() in content
    assert "Export generated on " in content
    # End-of-document footer closes the export.
    assert content.rstrip().endswith(wm.watermark_markdown_footer().rstrip())


def test_html_export_is_watermarked():
    from synapsis.exporters import export_html

    content, _ = export_html("Test", "abc123", _fake_rows(), "standard")
    assert 'class="ai-watermark"' in content
    assert wm.WATERMARK_BANNER in content
    # Per-page diagonal overlay + print footer, immediately after <body>.
    assert 'class="ai-watermark-overlay"' in content
    assert 'class="ai-watermark-pagefooter"' in content
    assert content.split("<body>", 1)[1].lstrip().startswith(
        '<div class="ai-watermark-overlay"'
    )
    assert ".ai-watermark-overlay{" in content  # CSS is embedded
    assert wm.contact_line_text().split(" —")[0] in content


def test_docx_export_is_watermarked(tmp_path):
    docx = pytest.importorskip("docx")
    from synapsis.exporters import export_docx

    filepath, _ = export_docx("My Session", "abc123", _fake_rows(), "standard", tmp_path)
    doc = docx.Document(filepath)
    assert doc.paragraphs[0].text == wm.WATERMARK_BANNER
    # The session title is wired into the running header.
    header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "My Session" in header_text
    assert wm.WATERMARK_BANNER in header_text
    # Diagonal per-page watermark shipped in the real export, not just in unit tests.
    assert "PowerPlusWaterMarkObject1" in doc.sections[0].header._element.xml
    assert "PAGE" in doc.sections[0].footer._element.xml


def test_html_dashboard_is_watermarked():
    from synapsis.exporters import render_dashboard_html

    html = render_dashboard_html(
        "Portfolio", [{"type": "kpi", "cards": [{"label": "Total", "value": "2,755"}]}]
    )
    assert 'class="ai-watermark"' in html
    assert wm.WATERMARK_BANNER in html
    assert 'class="ai-watermark-overlay"' in html
    assert 'class="ai-watermark-pagefooter"' in html


def test_workflow_run_exports_are_watermarked():
    from synapsis.exporters.workflow_run import (
        export_workflow_run_html,
        export_workflow_run_markdown,
    )

    run_log = {
        "workflow_name": "Portfolio scan",
        "run_id": "abcdef123456",
        "status": "completed",
        "initial_prompt": "Scan the portfolio.",
        "steps": [],
    }

    md, _ = export_workflow_run_markdown(run_log)
    assert wm.WATERMARK_BANNER in md
    assert wm.contact_line_text() in md
    assert md.rstrip().endswith(wm.watermark_markdown_footer().rstrip())

    html, _ = export_workflow_run_html(run_log)
    assert 'class="ai-watermark-overlay"' in html
    assert 'class="ai-watermark-pagefooter"' in html
    assert html.split("<body>", 1)[1].lstrip().startswith(
        '<div class="ai-watermark-overlay"'
    )
