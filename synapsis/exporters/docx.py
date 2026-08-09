"""
DOCX exporter — produces a Word document (.docx) from session rows.

Requires python-docx:  pip install python-docx
"""

import json
from datetime import datetime
from pathlib import Path

from fastapi import HTTPException

from .common import parse_row, safe_filename
from .watermark import apply_ai_watermark


def export_docx(title: str, session_id: str, rows, detail: str, export_dir: Path) -> tuple[str, str]:
    """Convert messages to a Word document and save it to export_dir.

    Args:
        title:      Human-readable session title.
        session_id: Session UUID embedded in the document metadata.
        rows:       Iterable of raw DB message rows (type, data JSON, ts).
        detail:     'standard' or 'full'. 'full' includes thinking blocks,
                    tool inputs/outputs, and non-file-upload system messages.
        export_dir: Directory in which the .docx file will be written.

    Returns:
        (filepath, filename) — absolute path string and the bare filename.

    Raises:
        HTTPException(500) if python-docx is not installed.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Install with: pip install python-docx")

    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(title, level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    # Metadata line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(
        f"Session: {session_id} · Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()

    for row in rows:
        msg_type, data, ts = parse_row(row)

        if msg_type == "user":
            heading = doc.add_heading(f"You ({ts})", level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            para = doc.add_paragraph(data.get("content", ""))
            para.style.font.size = Pt(11)

        elif msg_type == "text":
            heading = doc.add_heading(f"Assistant ({ts})", level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            for para_text in data.get("content", "").split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.style.font.size = Pt(11)

        elif msg_type == "system":
            subtype = data.get("subtype", "")
            content = data.get("content", "")
            if subtype == "file_upload":
                p = doc.add_paragraph()
                run = p.add_run(f"📎 {content}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                run.italic = True
            elif detail == "full":
                p = doc.add_paragraph()
                run = p.add_run(f"📋 System [{subtype}]: {content}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True

        elif msg_type == "thinking" and detail == "full":
            p = doc.add_paragraph()
            run = p.add_run(f"💭 Thinking: {data.get('content', '')}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
            run.italic = True

        elif msg_type == "tool_use":
            tool_name = data.get("tool", "unknown")
            if detail == "full":
                p = doc.add_paragraph()
                run = p.add_run(f"🔧 Tool: {tool_name}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
                run.bold = True
                p2 = doc.add_paragraph(json.dumps(data.get("input", {}), indent=2))
                p2.style.font.size = Pt(8)
                for r in p2.runs:
                    r.font.color.rgb = RGBColor(0x66, 0x33, 0x00)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"🔧 Tool: {tool_name}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
                run.italic = True

        elif msg_type == "tool_result" and detail == "full":
            content = str(data.get("content", ""))[:2000]
            p = doc.add_paragraph()
            run = p.add_run(f"📤 Result: {content}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x28, 0x35, 0x93)

        elif msg_type == "result":
            turns = data.get("turns", 0)
            duration = data.get("duration_ms", 0)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"— {turns} turns · {duration/1000:.1f}s —")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # AI-content watermark / disclaimer — required on every export. Applied
    # last so the notice box lands at the very top of the document body and
    # the per-page header/footer marks (running banner, diagonal draft
    # watermark, page numbers) are attached to the section.
    apply_ai_watermark(doc, title=title)

    export_dir.mkdir(parents=True, exist_ok=True)
    filename = safe_filename(title, session_id, "docx")
    filepath = export_dir / filename
    doc.save(str(filepath))
    return str(filepath), filename
